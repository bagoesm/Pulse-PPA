import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_doc_default_font_arial(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

def clear_body(doc):
    for p in list(doc.paragraphs):
        p._p.getparent().remove(p._p)
    for t in list(doc.tables):
        t._element.getparent().remove(t._element)

def set_cell_margins(cell, top=40, bottom=40, left=60, right=60):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_border_custom(cell, top=None, bottom=None, left=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    else:
        tcBorders.clear()

    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for b_name, val in borders.items():
        if val is not None:
            node = OxmlElement(f'w:{b_name}')
            if val:
                node.set(qn('w:val'), 'single')
                node.set(qn('w:sz'), '4') # 0.5pt
                node.set(qn('w:space'), '0')
                node.set(qn('w:color'), '000000')
            else:
                node.set(qn('w:val'), 'none')
            tcBorders.append(node)

def apply_font_arial(container, font_size=10.0, bold=False):
    if hasattr(container, 'rows'):
        for row in container.rows:
            for cell in row.cells:
                set_cell_margins(cell, top=40, bottom=40, left=60, right=60)
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.15
                    for r in p.runs:
                        r.font.name = 'Arial'
                        r.font.size = Pt(font_size)
                        if bold:
                            r.bold = True
    elif hasattr(container, 'paragraphs'):
        for p in container.paragraphs:
            p.paragraph_format.line_spacing = 1.15
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.size = Pt(font_size)
                if bold:
                    r.bold = True

def build_kop_header_table(header, doc_type):
    for p in list(header.paragraphs):
        p._p.getparent().remove(p._p)

    tbl = header.add_table(rows=1, cols=3, width=Inches(7.2))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [Inches(0.65), Inches(4.35), Inches(2.2)]

    row = tbl.rows[0]
    cell_logo = row.cells[0]
    cell_mid = row.cells[1]
    cell_right = row.cells[2]

    cell_logo.width = widths[0]
    cell_mid.width = widths[1]
    cell_right.width = widths[2]

    # 1. Logo KPPPA (Smaller)
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(0)
    run_logo = p_logo.add_run()
    run_logo.add_picture('public/LogoKPPPA.png', width=Inches(0.55))

    # 2. Ministry Text (Center - Title in ONE line)
    p_mid = cell_mid.paragraphs[0]
    p_mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_mid.paragraph_format.space_before = Pt(0)
    p_mid.paragraph_format.space_after = Pt(0)
    p_mid.paragraph_format.line_spacing = 1.05

    r1 = p_mid.add_run("KEMENTERIAN PEMBERDAYAAN PEREMPUAN DAN PERLINDUNGAN ANAK\nREPUBLIK INDONESIA\n")
    r1.bold = True
    r1.font.name = 'Arial'
    r1.font.size = Pt(8.0)

    r2 = p_mid.add_run("SEKRETARIAT KEMENTERIAN\n")
    r2.bold = True
    r2.font.name = 'Arial'
    r2.font.size = Pt(9.0)

    r3 = p_mid.add_run("JALAN MEDAN MERDEKA BARAT NOMOR 15 JAKARTA 10110\nTELEPON (021) 3842638, 3805563\nLaman: https://www.kemenpppa.go.id - Email: persuratan@kemenpppa.go.id")
    r3.font.name = 'Arial'
    r3.font.size = Pt(6.5)

    # 3. Right Subheader Text
    p_r = cell_right.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_r.paragraph_format.space_before = Pt(0)
    p_r.paragraph_format.space_after = Pt(0)
    p_r.paragraph_format.line_spacing = 1.05

    if doc_type == 'kwitansi':
        r_t1 = p_r.add_run("Tanda Pengeluaran\n")
        r_t1.bold = True
        r_t1.font.name = 'Arial'
        r_t1.font.size = Pt(9.0)

        r_t2 = p_r.add_run("Mata Anggaran :   {{mata_anggaran}}\nDibayarkan Tgl :   {{tanggal_dibayarkan}}\nPembukuan No :   {{pembukuan_no}}")
        r_t2.font.name = 'Arial'
        r_t2.font.size = Pt(7.5)
    elif doc_type == 'rincian':
        r_t1 = p_r.add_run("LAMPIRAN II\n")
        r_t1.bold = True
        r_t1.font.name = 'Arial'
        r_t1.font.size = Pt(8.5)

        r_t2 = p_r.add_run("PERATURAN MENTERI KEUANGAN NO 113/PMK.05/2012 TENTANG PERJALANAN DINAS JABATAN DALAM NEGERI BAGI PEJABAT NEGARA, PEGAWAI NEGERI DAN PEGAWAI TIDAK TETAP.")
        r_t2.bold = True
        r_t2.font.name = 'Arial'
        r_t2.font.size = Pt(6.5)
    elif doc_type == 'riil':
        r_t1 = p_r.add_run("LAMPIRAN II\n")
        r_t1.bold = True
        r_t1.font.name = 'Arial'
        r_t1.font.size = Pt(8.5)

        r_t2 = p_r.add_run("PERATURAN DIREKTUR JENDERAL PERBENDAHARAAN NOMOR PER ...../PB/2007 TENTANG PETUNJUK PELAKSANAAN DINAS JABATAN DALAM NEGERI BAGI PEJABAT NEGARA, PEGAWAI NEGERI DAN PEGAWAI TIDAK TETEP.")
        r_t2.bold = True
        r_t2.font.name = 'Arial'
        r_t2.font.size = Pt(6.5)

    for c in row.cells:
        set_cell_border_custom(c, top=False, bottom=False, left=False, right=False)

    p_line = header.add_paragraph()
    p_line.paragraph_format.space_before = Pt(2)
    p_line.paragraph_format.space_after = Pt(0)
    p_line_border = parse_xml(
        f'<w:pBdr {nsdecls("w")}>\n'
        f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="2E3B4E"/>\n'
        f'</w:pBdr>'
    )
    p_line._p.get_or_add_pPr().append(p_line_border)

def add_kwitansi_body(doc):
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(18)
    p_title.paragraph_format.space_after = Pt(12)
    rt = p_title.add_run("KWITANSI")
    rt.bold = True
    rt.font.name = 'Arial'
    rt.font.size = Pt(12)

    # Key Value Table (Rows 0-2: 3-cols, Row 3-5: Merged full width, Row 6: 3-cols)
    t1 = doc.add_table(rows=7, cols=3)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Inches(1.6), Inches(0.15), Inches(5.45)]

    # Rows 0-2
    kv1 = [
        ("Sudah Terima Dari", ":", "KUASA PENGGUNA ANGGARAN / PEJABAT PEMBUAT KOMITMEN\nKEMENTERIAN PEMBERDAYAAN PEREMPUAN DAN PERLINDUNGAN ANAK"),
        ("Uang Sebesar", ":", "Rp. {{uang_sebesar_format}}\n({{uang_sebesar_terbilang}})"),
        ("Untuk Pembayaran", ":", "Biaya perjalanan dinas menurut surat perjalanan dinas dari\nKEMENTERIAN PEMBERDAYAAN PEREMPUAN DAN PERLINDUNGAN ANAK\ntgl {{tanggal_spd}}\nNo. {{nomor_spd}}")
    ]

    for ri, row_data in enumerate(kv1):
        row = t1.rows[ri]
        for ci in range(3):
            cell = row.cells[ci]
            cell.width = widths[ci]
            p = cell.paragraphs[0]
            p.text = row_data[ci]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            set_cell_border_custom(cell, top=False, bottom=False, left=False, right=False)

    # Row 3: Maksud Perjalanan Dinas (Merged, Bold, No Colon)
    r3 = t1.rows[3]
    r3.cells[0].merge(r3.cells[2])
    p3 = r3.cells[0].paragraphs[0]
    r3_run = p3.add_run("Maksud Perjalanan Dinas")
    r3_run.bold = True
    p3.paragraph_format.space_before = Pt(4)
    p3.paragraph_format.space_after = Pt(2)

    # Row 4: Maksud Perjalanan Dinas Value (Merged, Normal font)
    r4 = t1.rows[4]
    r4.cells[0].merge(r4.cells[2])
    p4 = r4.cells[0].paragraphs[0]
    p4.text = "{{maksud_perjalanan_dinas}}"
    p4.paragraph_format.space_before = Pt(0)
    p4.paragraph_format.space_after = Pt(8)

    # Row 5: Pada... (Merged, Normal font)
    r5 = t1.rows[5]
    r5.cells[0].merge(r5.cells[2])
    p5 = r5.cells[0].paragraphs[0]
    p5.text = "Pada {{periode_perjadin}}"
    p5.paragraph_format.space_before = Pt(0)
    p5.paragraph_format.space_after = Pt(4)

    # Row 6: Jumlah (3-cols)
    r6 = t1.rows[6]
    r6_data = ("Jumlah", ":", "Rp. {{uang_sebesar_format}}")
    for ci in range(3):
        cell = r6.cells[ci]
        cell.width = widths[ci]
        p = cell.paragraphs[0]
        p.text = r6_data[ci]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        set_cell_border_custom(cell, top=False, bottom=False, left=False, right=False)

    apply_font_arial(t1, 10.0)

    # Generous gap paragraph before top signature table
    p_gap_sig = doc.add_paragraph()
    p_gap_sig.paragraph_format.space_before = Pt(24)
    p_gap_sig.paragraph_format.space_after = Pt(0)

    # Top Signatures Table (2 columns: Bendahara on Left, Yang Bertugas on Right)
    t_sig_top = doc.add_table(rows=1, cols=2)
    t_sig_top.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_top_widths = [Inches(5.1), Inches(2.1)]

    r0 = t_sig_top.rows[0]
    r0.cells[0].width = sig_top_widths[0]
    r0.cells[0].paragraphs[0].text = "Bendahara Pengeluaran\n\n\n\n\n{{nama_bendahara}}\nNIP. {{nip_bendahara}}"
    set_cell_border_custom(r0.cells[0], top=False, bottom=False, left=False, right=False)

    r0.cells[1].width = sig_top_widths[1]
    r0.cells[1].paragraphs[0].text = "Jakarta, {{tanggal_dibayarkan}}\nYang Bertugas,\n\n\n\n\n{{nama_pegawai}}\nNIP. {{nip_pegawai}}"
    set_cell_border_custom(r0.cells[1], top=False, bottom=False, left=False, right=False)

    apply_font_arial(t_sig_top, 9.5)

    # Gap paragraph to place PPK BELOW the upper NIPs ("lebih ke bawah")
    p_gap_ppk = doc.add_paragraph()
    p_gap_ppk.paragraph_format.space_before = Pt(20)
    p_gap_ppk.paragraph_format.space_after = Pt(0)

    # Bottom Signature Table (PPK - 1 Centered Column BELOW upper signatures)
    t_sig_ppk = doc.add_table(rows=1, cols=1)
    t_sig_ppk.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_ppk = t_sig_ppk.rows[0].cells[0]
    cell_ppk.width = Inches(4.5)

    p_ppk = cell_ppk.paragraphs[0]
    p_ppk.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ppk.text = "Mengetahui,\nA/n Kuasa Pengguna Anggaran,\nPejabat Pembuat Komitmen,\n\n\n\n{{nama_ppk}}\nNIP. {{nip_ppk}}"
    set_cell_border_custom(cell_ppk, top=False, bottom=False, left=False, right=False)

    apply_font_arial(t_sig_ppk, 9.5)

def add_rincian_body(doc):
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(18)
    p_title.paragraph_format.space_after = Pt(10)
    rt = p_title.add_run("RINCIAN BIAYA PERJALANAN DINAS")
    rt.bold = True
    rt.font.name = 'Arial'
    rt.font.size = Pt(11)

    # Key Value info
    t_kv = doc.add_table(rows=3, cols=3)
    widths_kv = [Inches(1.6), Inches(0.15), Inches(5.45)]
    kv_data = [
        ("Lampiran SPD NO.", ":", "{{nomor_spd}}"),
        ("Tanggal", ":", "{{tanggal_spd}}"),
        ("Maksud Perjalanan Dinas", ":", "{{maksud_perjalanan_dinas}}")
    ]
    for ri, rdata in enumerate(kv_data):
        row = t_kv.rows[ri]
        for ci in range(3):
            c = row.cells[ci]
            c.width = widths_kv[ci]
            p = c.paragraphs[0]
            p.text = rdata[ci]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(3)
            set_cell_border_custom(c, top=False, bottom=False, left=False, right=False)

    apply_font_arial(t_kv, 9.5)

    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(0)
    p_sp.paragraph_format.space_after = Pt(6)

    # Cost Table
    t_cost = doc.add_table(rows=13, cols=5)
    t_cost.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_widths = [Inches(0.45), Inches(2.55), Inches(1.85), Inches(1.35), Inches(1.0)]

    headers = ['No', 'PERINCIAN BIAYA', 'DETAIL', 'JUMLAH', 'KETERANGAN']
    h_row = t_cost.rows[0]
    for ci in range(5):
        c = h_row.cells[ci]
        c.width = c_widths[ci]
        p = c.paragraphs[0]
        p.text = headers[ci]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci in [0, 4] else (WD_ALIGN_PARAGRAPH.RIGHT if ci in [2, 3] else WD_ALIGN_PARAGRAPH.LEFT)

    cost_items = [
        ('1.', 'Uang Harian', '{{hari_harian}} Hari @ Rp. {{tarif_harian_format}}', 'Rp. {{total_harian_format}}', ''),
        ('2.', 'Uang Penginapan', '{{hari_penginapan}} Hari @ Rp. {{tarif_penginapan_format}}', 'Rp. {{total_penginapan_format}}', ''),
        ('3.', 'Transport - {{asal_transport1}} PP', '{{tujuan_transport1}} PP', 'Rp. {{total_transport1_format}}', ''),
        ('4.', 'Transport - {{asal_transport2}} PP', '{{tujuan_transport2}} PP', 'Rp. {{total_transport2_format}}', ''),
        ('5.', 'Uang Representatif', '{{hari_representatif}} Hari @ Rp. {{tarif_representatif_format}}', 'Rp. {{total_representatif_format}}', ''),
        ('6.', 'Uang Airport Tax', 'PP', 'Rp. {{total_airport_tax_format}}', ''),
        ('7.', 'Uang Transport Kantor - B/S/T (Taksi)', 'PP', 'Rp. {{total_transport_taksi_format}}', ''),
        ('8.', 'Uang Transport B/S/T - Lokasi (Taksi)', 'PP', 'Rp. {{total_transport_lokasi_format}}', ''),
        ('9.', 'Uang Transport Kota/Kab.', 'PP', 'Rp. {{total_transport_kota_format}}', ''),
        ('10.', 'Sewa Kendaraan (Roda 4)', '', 'Rp. {{total_sewa_kendaraan_format}}', ''),
    ]

    for ri, item in enumerate(cost_items, start=1):
        row = t_cost.rows[ri]
        for ci in range(5):
            c = row.cells[ci]
            c.width = c_widths[ci]
            p = c.paragraphs[0]
            p.text = item[ci]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if ci == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif ci in [2, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Row 11 & 12 Footer Merged
    r11 = t_cost.rows[11]
    r11.cells[0].merge(r11.cells[4])
    p11 = r11.cells[0].paragraphs[0]
    p11.text = "Jumlah Yang Dibayarkan : Rp. {{uang_sebesar_format}}"
    p11.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r12 = t_cost.rows[12]
    r12.cells[0].merge(r12.cells[4])
    p12 = r12.cells[0].paragraphs[0]
    p12.text = "Terbilang : {{uang_sebesar_terbilang}}"
    p12.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Apply exact borders (Outer left border on col 0 IS ALWAYS TRUE!)
    for ri, row in enumerate(t_cost.rows):
        for ci, cell in enumerate(row.cells):
            if ri == 0:
                set_cell_border_custom(cell, top=True, bottom=True, left=True, right=True)
            elif ri in [11, 12]:
                set_cell_border_custom(cell, top=True, bottom=True, left=True, right=True)
            else:
                set_cell_border_custom(cell, top=False, bottom=False, left=(ci == 0), right=True)

    apply_font_arial(t_cost, 9.0)

    # Gap paragraph to separate from signature table
    p_gap_sig = doc.add_paragraph()
    p_gap_sig.paragraph_format.space_before = Pt(20)
    p_gap_sig.paragraph_format.space_after = Pt(0)

    # Top Signatures Table (2 columns: Bendahara on Left, Yang Menerima on Right)
    t_sig_top = doc.add_table(rows=1, cols=2)
    t_sig_top.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_top_widths = [Inches(5.1), Inches(2.1)]

    r0 = t_sig_top.rows[0]
    r0.cells[0].width = sig_top_widths[0]
    r0.cells[0].paragraphs[0].text = "Telah Dibayar Sejumlah\nRp. {{uang_sebesar_format}}\nBendahara Pengeluaran\n\n\n\n\n{{nama_bendahara}}\nNIP. {{nip_bendahara}}"
    set_cell_border_custom(r0.cells[0], top=False, bottom=False, left=False, right=False)

    r0.cells[1].width = sig_top_widths[1]
    r0.cells[1].paragraphs[0].text = "Jakarta, {{tanggal_dibayarkan}}\nTelah menerima jumlah uang\nRp. {{uang_sebesar_format}}\nYang Menerima,\n\n\n\n{{nama_pegawai}}\nNIP. {{nip_pegawai}}"
    set_cell_border_custom(r0.cells[1], top=False, bottom=False, left=False, right=False)

    apply_font_arial(t_sig_top, 9.0)

    # Gap paragraph to place PPK BELOW the upper NIPs ("lebih ke bawah")
    p_gap_ppk = doc.add_paragraph()
    p_gap_ppk.paragraph_format.space_before = Pt(20)
    p_gap_ppk.paragraph_format.space_after = Pt(0)

    # Bottom Signature Table (PPK - 1 Centered Column BELOW upper signatures)
    t_sig_ppk = doc.add_table(rows=1, cols=1)
    t_sig_ppk.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_ppk = t_sig_ppk.rows[0].cells[0]
    cell_ppk.width = Inches(4.5)

    p_ppk = cell_ppk.paragraphs[0]
    p_ppk.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ppk.text = "Mengetahui,\nPejabat Pembuat Komitmen,\n\n\n\n{{nama_ppk}}\nNIP. {{nip_ppk}}"
    set_cell_border_custom(cell_ppk, top=False, bottom=False, left=False, right=False)

    apply_font_arial(t_sig_ppk, 9.0)

def add_riil_body(doc):
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(18)
    p_title.paragraph_format.space_after = Pt(10)
    rt = p_title.add_run("DAFTAR PENGELUARAN RIIL")
    rt.bold = True
    rt.font.name = 'Arial'
    rt.font.size = Pt(11)

    # Key Value info
    t_kv = doc.add_table(rows=3, cols=3)
    widths_kv = [Inches(1.6), Inches(0.15), Inches(5.45)]
    kv_data = [
        ("Nama", ":", "{{nama_pegawai}}"),
        ("NIP", ":", "{{nip_pegawai}}"),
        ("Maksud Perjalanan Dinas", ":", "{{maksud_perjalanan_dinas}}")
    ]
    for ri, rdata in enumerate(kv_data):
        row = t_kv.rows[ri]
        for ci in range(3):
            c = row.cells[ci]
            c.width = widths_kv[ci]
            p = c.paragraphs[0]
            p.text = rdata[ci]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(3)
            set_cell_border_custom(c, top=False, bottom=False, left=False, right=False)

    apply_font_arial(t_kv, 9.5)

    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(0)
    p_sp.paragraph_format.space_after = Pt(6)

    # Statement 1
    p_st1 = doc.add_paragraph()
    p_st1.paragraph_format.space_before = Pt(4)
    p_st1.paragraph_format.space_after = Pt(6)
    p_st1.paragraph_format.line_spacing = 1.15

    r_st1_a = p_st1.add_run("Berdasar Surat Perjalanan Dinas (SPD) tanggal {{tanggal_spd}} Nomor. {{nomor_spd}} dengan ini kami menyatakan dengan sesungguhnya bahwa :\n")
    r_st1_a.font.name = 'Arial'
    r_st1_a.font.size = Pt(9.5)

    r_st1_b = p_st1.add_run("1. Biaya transport pegawai dan/atau biaya penginapan dibawah ini yang tidak dapat diperoleh bukti bukti pengeluaranya meliputi :")
    r_st1_b.font.name = 'Arial'
    r_st1_b.font.size = Pt(9.5)

    # Cost Table (3 columns)
    t_riil = doc.add_table(rows=8, cols=3)
    t_riil.alignment = WD_TABLE_ALIGNMENT.CENTER
    r_widths = [Inches(0.6), Inches(4.9), Inches(1.7)]

    h_row = t_riil.rows[0]
    headers = ['No', 'URAIAN', 'JUMLAH']
    for ci in range(3):
        c = h_row.cells[ci]
        c.width = r_widths[ci]
        p = c.paragraphs[0]
        p.text = headers[ci]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else (WD_ALIGN_PARAGRAPH.RIGHT if ci == 2 else WD_ALIGN_PARAGRAPH.LEFT)

    riil_items = [
        ('A.', 'Transport Kantor - B/S/T PP', 'Rp. {{riil_transport_taksi_format}}'),
        ('B.', 'Transport B/S/T - Lokasi PP', 'Rp. {{riil_transport_lokasi_format}}'),
        ('C.', 'Uang Transport Kota/Kab. (Taksi) PP', 'Rp. {{riil_transport_kota_format}}'),
        ('D.', '', 'Rp. 0'),
        ('E.', '', 'Rp. 0'),
        ('F.', '', 'Rp. 0')
    ]

    for ri, item in enumerate(riil_items, start=1):
        row = t_riil.rows[ri]
        for ci in range(3):
            c = row.cells[ci]
            c.width = r_widths[ci]
            p = c.paragraphs[0]
            p.text = item[ci]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if ci == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif ci == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Row 7 Footer
    r7 = t_riil.rows[7]
    r7.cells[0].paragraphs[0].text = ""
    r7.cells[1].paragraphs[0].text = "Jumlah Yang Dibayarkan :"
    r7.cells[2].paragraphs[0].text = "Rp. {{riil_total_format}}"
    r7.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # FIX OUTER LEFT BORDER (left=(ci == 0) ensures leftmost column border is ALWAYS drawn!)
    for ri, row in enumerate(t_riil.rows):
        for ci, cell in enumerate(row.cells):
            if ri == 0:
                set_cell_border_custom(cell, top=True, bottom=True, left=True, right=True)
            elif ri == 7:
                set_cell_border_custom(cell, top=True, bottom=True, left=True, right=True)
            else:
                set_cell_border_custom(cell, top=False, bottom=False, left=(ci == 0), right=True)

    apply_font_arial(t_riil, 9.5)

    p_sp2 = doc.add_paragraph()
    p_sp2.paragraph_format.space_before = Pt(0)
    p_sp2.paragraph_format.space_after = Pt(6)

    # Statement 2
    p_st2 = doc.add_paragraph()
    p_st2.paragraph_format.space_before = Pt(6)
    p_st2.paragraph_format.space_after = Pt(8)
    p_st2.paragraph_format.line_spacing = 1.15

    r_st2_a = p_st2.add_run("2. Jumlah uang tersebut pada angka 1 benar benar dikeluarkan untuk pelaksanaan perjalanan dinas dimaksud dan apabila dikemudian hari terdapat kelebihan atas pembayaran, kami bersedia untuk menyetorkan kelebihan tersebut ke Kas Negara.\n")
    r_st2_a.font.name = 'Arial'
    r_st2_a.font.size = Pt(9.5)

    r_st2_b = p_st2.add_run("Demikian Pernyataan ini kami buat dengan sebenar benarnya, untuk digunakan sebagai mana mestinya.")
    r_st2_b.font.name = 'Arial'
    r_st2_b.font.size = Pt(9.5)

    # Signatures
    t_sig = doc.add_table(rows=1, cols=2)
    sig_cells = [
        "Bendahara Pengeluaran\n\n\n\n{{nama_bendahara}}\nNIP. {{nip_bendahara}}",
        "Jakarta, {{tanggal_dibayarkan}}\nPejabat Negara/Pegawai Negeri yang melakukan perjalanan dinas\n\n\n\n{{nama_pegawai}}\nNIP. {{nip_pegawai}}"
    ]
    for ci in range(2):
        c = t_sig.rows[0].cells[ci]
        p = c.paragraphs[0]
        p.text = sig_cells[ci]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        set_cell_border_custom(c, top=False, bottom=False, left=False, right=False)

    apply_font_arial(t_sig, 9.5)

def build_kwitansi():
    doc = docx.Document('public/surat-keterangan.docx')
    set_doc_default_font_arial(doc)
    clear_body(doc)
    s = doc.sections[0]
    s.top_margin = Inches(0.25)
    s.bottom_margin = Inches(0.4)
    s.left_margin = Inches(0.5)
    s.right_margin = Inches(0.5)
    s.header_distance = Inches(0.1)

    build_kop_header_table(s.header, 'kwitansi')
    add_kwitansi_body(doc)
    doc.save('public/simperjadin-kwitansi.docx')
    print("Created public/simperjadin-kwitansi.docx")

def build_rincian():
    doc = docx.Document('public/surat-keterangan.docx')
    set_doc_default_font_arial(doc)
    clear_body(doc)
    s = doc.sections[0]
    s.top_margin = Inches(0.25)
    s.bottom_margin = Inches(0.4)
    s.left_margin = Inches(0.5)
    s.right_margin = Inches(0.5)
    s.header_distance = Inches(0.1)

    build_kop_header_table(s.header, 'rincian')
    add_rincian_body(doc)
    doc.save('public/simperjadin-rincian.docx')
    print("Created public/simperjadin-rincian.docx")

def build_riil():
    doc = docx.Document('public/surat-keterangan.docx')
    set_doc_default_font_arial(doc)
    clear_body(doc)
    s = doc.sections[0]
    s.top_margin = Inches(0.25)
    s.bottom_margin = Inches(0.4)
    s.left_margin = Inches(0.5)
    s.right_margin = Inches(0.5)
    s.header_distance = Inches(0.1)

    build_kop_header_table(s.header, 'riil')
    add_riil_body(doc)
    doc.save('public/simperjadin-riil.docx')
    print("Created public/simperjadin-riil.docx")

def build_lengkap():
    doc = docx.Document('public/surat-keterangan.docx')
    set_doc_default_font_arial(doc)
    clear_body(doc)

    # Section 1: Kwitansi
    s1 = doc.sections[0]
    s1.top_margin = Inches(0.25)
    s1.bottom_margin = Inches(0.4)
    s1.left_margin = Inches(0.5)
    s1.right_margin = Inches(0.5)
    s1.header_distance = Inches(0.1)
    build_kop_header_table(s1.header, 'kwitansi')
    add_kwitansi_body(doc)

    # Section 2: Rincian
    s2 = doc.add_section(WD_SECTION.NEW_PAGE)
    s2.top_margin = Inches(0.25)
    s2.bottom_margin = Inches(0.4)
    s2.left_margin = Inches(0.5)
    s2.right_margin = Inches(0.5)
    s2.header_distance = Inches(0.1)
    s2.header.is_linked_to_previous = False
    build_kop_header_table(s2.header, 'rincian')
    add_rincian_body(doc)

    # Section 3: Riil
    s3 = doc.add_section(WD_SECTION.NEW_PAGE)
    s3.top_margin = Inches(0.25)
    s3.bottom_margin = Inches(0.4)
    s3.left_margin = Inches(0.5)
    s3.right_margin = Inches(0.5)
    s3.header_distance = Inches(0.1)
    s3.header.is_linked_to_previous = False
    build_kop_header_table(s3.header, 'riil')
    add_riil_body(doc)

    doc.save('public/simperjadin-lengkap.docx')
    print("Created public/simperjadin-lengkap.docx")

if __name__ == '__main__':
    build_kwitansi()
    build_rincian()
    build_riil()
    build_lengkap()
